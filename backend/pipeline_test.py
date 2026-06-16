import requests
import os
import re
import datetime
from pathlib import Path
from openai import OpenAI

# ── Fill these in locally ────────────────────────────────────────────────────
SARVAM_KEY = "your_sarvam_key_here"
OPENAI_KEY = "your_openai_key_here"

# ── Sarvam transcription ─────────────────────────────────────────────────────
def transcribe(audio_path):
    print(f"[1/3] Transcribing with Sarvam: {Path(audio_path).name}")
    with open(audio_path, "rb") as f:
        resp = requests.post(
            "https://api.sarvam.ai/speech-to-text",
            headers={"api-subscription-key": SARVAM_KEY},
            files={"file": (Path(audio_path).name, f, "audio/wav")},
            data={"language_code": "hi-IN", "model": "saaras:v3", "with_timestamps": "true", "with_diarization": "true"},
            timeout=300,
        )
        resp.raise_for_status()

    data = resp.json()
    segments, speaker_map, counter = [], {}, 1

    if "diarized_transcript" in data:
        for turn in data["diarized_transcript"]:
            spk = turn.get("speaker", "SPEAKER_0")
            if spk not in speaker_map:
                speaker_map[spk] = f"Speaker {counter}"; counter += 1
            segments.append({"speaker": speaker_map[spk], "start": turn.get("start", 0), "end": turn.get("end", 0), "text": turn.get("transcript", "").strip()})
    else:
        segments.append({"speaker": "Speaker 1", "start": 0, "end": 0, "text": data.get("transcript", "")})

    full_text = "\n".join(f"{s['speaker']} [{_fmt(s['start'])}-{_fmt(s['end'])}]: {s['text']}" for s in segments)
    print(f"      ✓ {len(segments)} segments")
    return full_text, segments

def _fmt(s):
    try:
        m, sec = divmod(int(float(s)), 60); return f"{m}:{sec:02d}"
    except: return "0:00"

# ── GPT-4o-mini summary ───────────────────────────────────────────────────────
PROMPT = """You are an expert analyst for a fintech company onboarding MFDs (Mutual Fund Distributors) in India.

Analyse this Hinglish sales call transcript. Produce a detailed, specific summary in English using ONLY information present in the transcript.

## Participants
- Name, role, company for each speaker

## Call Context
- Type of call, platform discussed

## Prospect Background
- Current setup, AUM, years of experience, pain points

## Key Topics Discussed
- Each topic with specific detail

## Objections & Concerns Raised
- Each objection with context

## Positive Signals
- Specific quotes or moments showing interest

## Brokerage & Commercial Discussion
- Any numbers, percentages, commission details

## Action Points
| # | Action | Owner |
|---|--------|-------|

## Overall Assessment
- Conversion likelihood: High/Medium/Low
- Reasoning and recommended next step

TRANSCRIPT:
{transcript}"""

def summarise(transcript):
    print("[2/3] Summarising with GPT-4o-mini...")
    client = OpenAI(api_key=OPENAI_KEY)
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert business analyst specialising in Indian fintech and MFD onboarding."},
            {"role": "user", "content": PROMPT.format(transcript=transcript[:12000])}
        ],
        max_tokens=2000,
        temperature=0.3,
    )
    summary = resp.choices[0].message.content.strip()
    print("      ✓ Summary generated")
    return summary

# ── Build .docx ───────────────────────────────────────────────────────────────
COLORS = [(0x1F,0x4E,0x79),(0x83,0x31,0x22),(0x37,0x5C,0x23),(0x4A,0x23,0x6E)]

def build_docx(filename, segments, summary, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print(f"[3/3] Building report → {output_path}")
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1.2)

    def blue_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs: r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Call Transcription & Summary Report")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(f"{filename}   ·   {datetime.datetime.now().strftime('%d %b %Y, %H:%M')}").font.size = Pt(9)
    doc.add_paragraph()

    blue_heading("Call Summary")
    for line in summary.split("\n"):
        line = line.rstrip()
        if not line: doc.add_paragraph(); continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for r in p.runs: r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _bold(p, re.sub(r"^[-*]\s+","",line))
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c)<=set("- :") for c in cells): continue
            p = doc.add_paragraph(); p.add_run("   ".join(cells)).font.size = Pt(9)
        else:
            p = doc.add_paragraph(); _bold(p, line)

    doc.add_page_break()
    blue_heading("Full Transcript")
    spk_idx, prev = {}, None
    for seg in segments:
        spk = seg["speaker"]
        if spk not in spk_idx: spk_idx[spk] = len(spk_idx)
        if spk != prev:
            prev = spk
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            sr = p.add_run(f"{spk}  "); sr.bold = True
            sr.font.color.rgb = RGBColor(*COLORS[spk_idx[spk]%len(COLORS)])
            p.add_run(f"[{_fmt(seg['start'])}–{_fmt(seg['end'])}]  ").font.size = Pt(8)
            p.add_run(seg["text"]).font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(seg["text"]).font.size = Pt(10)
    doc.save(output_path)
    print(f"      ✓ Saved")

def _bold(p, text):
    for i, part in enumerate(re.split(r"\*\*(.*?)\*\*", text)):
        p.add_run(part).bold = (i%2==1)

# ── Run ───────────────────────────────────────────────────────────────────────
import sys
audio = sys.argv[1] if len(sys.argv) > 1 else "sample_call_clean.wav"
output = str(Path(audio).with_suffix("")) + "_report.docx"
transcript, segments = transcribe(audio)
summary = summarise(transcript)
build_docx(Path(audio).name, segments, summary, output)
print(f"\n✅ Done! Report saved to: {output}\n")
